import io

from bs4 import BeautifulSoup

from odoo import api, models, fields
from odoo.exceptions import UserError
from odoo.tools import ustr, groupby
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
import base64
import os
import tempfile
import subprocess
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


class SaleOrder(models.Model):
    _inherit = 'sale.order'

    def action_generate_offer_report(self):
        for line in self.order_line:
            partner_lang = self.partner_id.lang or 'en_US'
            offer_template = line.product_id.product_tmpl_id.offer_template_ids.filtered(
                lambda t: t.language_id.code == partner_lang)
            if not offer_template:
                raise UserError(
                    f'Please upload an offer template for the product {line.product_id.name} in language {partner_lang}.')

            offer_template = offer_template[0]  # Wybieramy pierwszy szablon, jeśli są wielokrotne

            # Download the template from the filestore
            template_path = self._get_file_from_filestore(offer_template.template_attachment_id)

            # Create a temporary file for the output docx
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
                output_path = temp_docx.name

            # Context with dynamic values from the sale order
            context = {
                "offer_number": self.name,
                "validity_date": self.validity_date.strftime('%d/%m/%Y') if self.validity_date else ' ',
                "partner_name": self.partner_id.name,
                "partner_street": self.partner_id.street,
                "partner_city": self.partner_id.city,
                "partner_zip": self.partner_id.zip,
                "partner_phone": self.partner_id.phone if self.partner_id.phone else self.partner_id.mobile if self.partner.mobile else ' ',
                "amount_untaxed": f"{self.amount_untaxed}",
                "amount_total": f"{self.amount_total}",
                "amount_tax": f"{self.amount_tax}",
                "salesperson_name": self.user_id.name,
                "salesperson_phone": self.user_id.phone if self.user_id.phone else self.user_id.mobile if self.user_id.mobile else ' ',
                "salesperson_email": self.user_id.email,
                "salesperson_company": self.user_id.parent_id.name,
                "salesperson_company_street": self.user_id.parent_id.street,
                "salesperson_company_city": self.user_id.parent_id.city,
                "salesperson_company_zip": self.user_id.parent_id.zip if self.user_id.parent_id.zip else ' ',
                "salesperson_www": self.user_id.parent_id.website,
                "dealer_logo": self.user_id.parent_id.image_1920,  # Assuming 'image_1920' stores the dealer logo
            }


            # Add options to context with individual placeholders for each PTAV
            options_context = self._get_options_context()
            context.update(options_context)

            # Log all context keys
            print("Context keys:")
            for key in context.keys():
                print(f'<<{key}>>')

            # Generate the report
            generator = OfferReportGenerator(template_path, output_path)
            pdf_path = generator.generate_report(context)

            # Save the generated PDF to the filestore and attach to the record
            attachment = self._save_file_to_filestore(pdf_path)

            # Clean up temporary files
            os.remove(output_path)
            # os.remove(pdf_path)

            # Return action to download the generated PDF
            return {
                'type': 'ir.actions.act_url',
                'url': f'/web/content/{attachment.id}?download=true',
                'target': 'self',
            }

    def _get_file_from_filestore(self, attachment):
        filestore = self.env['ir.attachment']._filestore()
        return os.path.join(filestore, attachment.store_fname)

    def _save_file_to_filestore(self, file_path):
        with open(file_path, 'rb') as file:
            data = file.read()

        attachment = self.env['ir.attachment'].create({
            'name': os.path.basename(file_path),
            'type': 'binary',
            'datas': base64.b64encode(data).decode('utf-8'),
            'res_model': self._name,
            'res_id': self.id,
            'mimetype': 'application/pdf'
        })
        return attachment

    def _get_options_context(self):
        options_context = {}
        for line in self.order_line:
            no_variant_ptavs = line.product_no_variant_attribute_value_ids._origin.filtered(
                # Only describe the attributes where a choice was made by the customer
                lambda ptav: ptav.display_type == 'multi' or ptav.attribute_line_id.value_count > 1
            )
            custom_ptavs = line.product_custom_attribute_value_ids.custom_product_template_attribute_value_id
            multi_ptavs = no_variant_ptavs.filtered(lambda ptav: ptav.display_type == 'multi').sorted()

            for ptav in multi_ptavs:
                ptav_key_title = f"option_{ptav.catalogue_number}_title"
                ptav_key_desc = f"option_{ptav.catalogue_number}_desc"
                options_context[ptav_key_title] = ptav.name
                options_context[ptav_key_desc] = self._convert_html_to_text(ptav.quotation_description or '')

        return options_context


    # def _get_options_text(self):
    #     options_text = ""
    #     for line in self.order_line:
    #         no_variant_ptavs = line.product_no_variant_attribute_value_ids._origin.filtered(
    #             # Only describe the attributes where a choice was made by the customer
    #             lambda ptav: ptav.display_type == 'multi' or ptav.attribute_line_id.value_count > 1
    #         )
    #         custom_ptavs = line.product_custom_attribute_value_ids.custom_product_template_attribute_value_id
    #         multi_ptavs = no_variant_ptavs.filtered(lambda ptav: ptav.display_type == 'multi').sorted()
    #
    #         for pta, ptavs in groupby(multi_ptavs, lambda ptav: ptav.attribute_id):
    #             for ptav in ptavs:
    #                 options_text += f"- {ptav.name}\n{self._convert_html_to_text(ptav.quotation_description or '')}\n\n"
    #
    #     print(f'options: {options_text}')
    #     return options_text

    def _convert_html_to_text(self, html_content):
        # Use BeautifulSoup to convert HTML content to plain text with basic formatting
        soup = BeautifulSoup(html_content, 'html.parser')
        text = ''
        for element in soup.descendants:
            if element.name == 'li':
                text += f'\u2022 {element.get_text(strip=True)}\n'
            elif element.name == 'br':
                text += '\n'
            elif element.name in ['p', 'div']:
                text += f'{element.get_text(strip=True)}\n\n'
            # elif isinstance(element, str):
            #     text += element
        return text.strip()


class OfferReportGenerator:
    def __init__(self, template_path, output_path):
        self.template_path = template_path
        self.output_path = output_path

    def generate_report(self, context):
        # Load the Word template
        doc = Document(self.template_path)

        # Replace placeholders with actual data from context while preserving formatting in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in context.items():
                if f'<<{key}>>' in paragraph.text:
                    self.replace_placeholder_in_paragraph(paragraph, f'<<{key}>>', str(value))

        # Replace placeholders in headers and footers
        for section in doc.sections:
            header = section.header
            footer = section.footer
            for paragraph in header.paragraphs + footer.paragraphs:
                for key, value in context.items():
                    if f'<<{key}>>' in paragraph.text:
                        self.replace_placeholder_in_paragraph(paragraph, f'<<{key}>>', str(value))

            # Replace placeholders in tables within headers and footers
            for table in header.tables + footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in context.items():
                                if f'<<{key}>>' in paragraph.text:
                                    self.replace_placeholder_in_paragraph(paragraph, f'<<{key}>>', str(value))

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in context.items():
                            if f'<<{key}>>' in paragraph.text:
                                self.replace_placeholder_in_paragraph(paragraph, f'<<{key}>>', str(value))

        # Replace image placeholders
        if 'dealer_logo' in context and context['dealer_logo']:
            self.replace_image(doc, '<<dealer_logo>>', context['dealer_logo'])

        # Save the modified document
        doc.save(self.output_path)

        pdf_path = self.output_path

        # Convert to PDF using LibreOffice
        # pdf_path = os.path.splitext(self.output_path)[0] + ".pdf"
        # self.convert_to_pdf(self.output_path, pdf_path)

        # Return the PDF path
        return pdf_path


    def replace_image(self, doc, placeholder, image_data):
        # Save the image data to a temporary file
        image_path = self._save_temp_image(image_data)

        # Iterate through all shapes in the document to find the image placeholder using alt_text
        print(f'shapes len: {len(doc.inline_shapes)}')
        for shape in doc.inline_shapes:
            descr = shape._inline.graphic.graphicData.pic.nvPicPr.cNvPr.get('descr')
            # print(f'00 shape descr: {descr}')
            if shape._inline.graphic.graphicData.pic.nvPicPr.cNvPr.get('descr') == placeholder:
                # Replace the existing image
                with open(image_path, 'rb') as img_file:
                    new_image_part = doc.part.related_parts[shape._inline.graphic.graphicData.pic.blipFill.blip.embed]
                    new_image_part._blob = img_file.read()

        # for section in doc.sections:
        #     for header_footer in [section.header, section.footer]:
        #         for shape in header_footer.shapes:
        #             if shape.shape_type == 13:  # 13 indicates a picture shape
        #                 if shape._element.xpath(".//a:blip"):
        #                     if shape._element.xpath(".//a:blip")[
        #                         0].getparent().getparent().get("descr") == placeholder:
        #                         new_blip = parse_xml

        # Replace images in headers and footers
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                for shape in header_footer._element.xpath(".//a:blip"):
                    # print(f'0a shape descr: {shape.get("descr")}')
                    # print(f'1a shape descr: {shape.getparent().get("descr")}')
                    # print(f'2a shape descr: {shape.getparent().getparent().get("descr")}')
                    # print(f'3a shape descr: {shape.getparent().getparent().getparent().get("descr")}')
                    if shape.getparent().getparent().getparent().get("descr") == placeholder:
                        new_blip = parse_xml(
                            f'<a:blip {nsdecls("a")} r:embed="{self._add_image_to_doc(doc, image_path)}"/>')
                        shape.getparent().replace(shape, new_blip)

                for table in header_footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for shape in cell._element.xpath(".//a:blip"):
                                # print(f'0 shape descr: {shape.get("descr")}')
                                # print(f'1 shape descr: {shape.getparent().get("descr")}')
                                # print(f'2 shape descr: {shape.getparent().getparent().get("descr")}')
                                # print(f'3 shape descr: {shape.getparent().getparent().getparent().get("descr")}')
                                if shape.getparent().getparent().getparent().get("descr") == placeholder:
                                    new_blip = parse_xml(
                                        f'<a:blip {nsdecls("a")} r:embed="{self._add_image_to_doc(doc, image_path)}"/>')
                                    shape.getparent().replace(shape, new_blip)


        os.remove(image_path)



    def _save_temp_image(self, image_data):
        # Save the image data to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as temp_image:
            temp_image.write(base64.b64decode(image_data))
            return temp_image.name

    def replace_placeholder_in_paragraph(self, paragraph, placeholder, value):
        # Concatenate all runs to a single string
        full_text = "".join(run.text for run in paragraph.runs)

        # Replace placeholder in the concatenated text
        updated_text = full_text.replace(placeholder, value)

        # Update the runs with new text while preserving formatting
        if updated_text != full_text:
            # Clear the existing runs
            for run in paragraph.runs:
                run.text = ""

            # Assign the updated text to the first run
            paragraph.runs[0].text = updated_text
    def convert_to_pdf(self, input_path, output_path):
        # Use LibreOffice in headless mode to convert .docx to .pdf
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(output_path),
                 input_path],
                check=True
            )
        except subprocess.CalledProcessError as e:
            raise UserError(f"Failed to convert to PDF: {str(e)}")
